# ==========================================
# server.py – Flask API backend for RoleMatch AI UI
# Bridges the HTML frontend with the Python logic
# from jobauto.py (Gemini, SMTP, CSV tracking)
# ==========================================

import os
import json
import re
import smtplib
import tempfile
from datetime import datetime
from email.message import EmailMessage
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv
import pandas as pd

# Google Sheets tracker (per-user)
try:
    from sheets_manager import get_or_create_sheet, log_sent_email, get_sheet_url, get_sheet_data, list_all_user_sheets
    SHEETS_AVAILABLE = True
except Exception as e:
    print(f"Warning: Google Sheets integration not available: {e}")
    SHEETS_AVAILABLE = False

# PDF reading
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX reading
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# HTTP for OpenRouter API
import requests as http_requests

load_dotenv()

app = Flask(__name__, static_folder='ui', static_url_path='')
CORS(app)

# ── Config ──
CSV_FILE = "job_tracker.csv"
UPLOAD_DIR = "/tmp/uploads" if os.environ.get("VERCEL") else os.path.abspath("uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ── Removed in-memory single-user session store (now stateless via frontend / local storage) ──

# =========================
# STATIC FILE SERVING (UI)
# =========================

@app.route('/')
def serve_login():
    return send_from_directory('ui', 'login.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('ui', path)


# =========================
# CSV TRACKER & SHEETS (Replaced by Supabase)
# =========================
def load_tracker_data():
    return pd.DataFrame()



# =========================
# RESUME TEXT EXTRACTION
# =========================

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    if not PDF_AVAILABLE:
        return ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except:
        return ""

def extract_name_from_text(text):
    return ""


# =========================
# API ROUTES
# =========================

# ── Health check ──
@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "openrouter_key_loaded": bool(os.getenv("OPENROUTER_API_KEY", ""))})


# ── Config endpoints ──
@app.route('/api/config', methods=['GET'])
def get_config():
    return jsonify({
        "supabase_url": os.getenv("SUPABASE_URL", ""),
        "supabase_key": os.getenv("SUPABASE_ANON_KEY", "")
    })


# ── Login (simple session) ──
@app.route('/api/login', methods=['POST'])
def login():
    data = request.json or {}
    email = data.get('email', '')
    password = data.get('password', '')
    full_name = data.get('full_name', '')
    if not email or not password:
        return jsonify({"success": False, "error": "Email and password are required"}), 400
    return jsonify({"success": True, "email": email, "sheet_url": None})

@app.route('/api/logout', methods=['POST'])
def logout():
    return jsonify({"success": True})


# ── Set sender credentials ──
@app.route('/api/credentials', methods=['POST'])
def set_credentials():
    data = request.json or {}
    email = data.get("email", "")
    full_name = data.get("full_name", "")
    # Session completely managed statelessly now. 
    sheet_url = None


    return jsonify({"success": True, "sheet_url": sheet_url})


# ── Upload resume ──
@app.route('/api/upload/resume', methods=['POST'])
def upload_resume():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({"success": False, "error": "Empty filename"}), 400

    import uuid
    safe_name = file.filename.replace(' ', '_')
    unique_id = str(uuid.uuid4())[:8]
    save_path = os.path.join(UPLOAD_DIR, f"resume_{unique_id}_{safe_name}")
    file.save(save_path)

    # Extract text from resume
    ext = file.filename.rsplit('.', 1)[-1].lower()
    resume_text = ""

    if ext == 'pdf':
        resume_text = extract_text_from_pdf(save_path)
    elif ext == 'txt':
        with open(save_path, 'r', encoding='utf-8', errors='ignore') as f:
            resume_text = f.read()
    elif ext in ['docx', 'doc'] and DOCX_AVAILABLE:
        try:
            doc = Document(save_path)
            resume_text = '\n'.join([p.text for p in doc.paragraphs])
        except:
            resume_text = ""

    resume_name = extract_name_from_text(resume_text)

    return jsonify({
        "success": True,
        "filename": file.filename,
        "size": os.path.getsize(save_path),
        "text_preview": resume_text[:500] if resume_text else "",
        "resume_text": resume_text,
        "resume_name": resume_name,
        "resume_path": save_path
    })


from werkzeug.utils import secure_filename

# ── Upload TXT scrape file ──
@app.route('/api/upload/txt', methods=['POST'])
def upload_txt():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({"success": False, "error": "Empty filename"}), 400

    content = file.read().decode('utf-8', errors='ignore')

    # Also save to scanned_jobs folder
    scanned_dir = "/tmp/scanned_jobs" if os.environ.get("VERCEL") else "scanned_jobs"
    os.makedirs(scanned_dir, exist_ok=True)
    
    # Strip any paths to fix weird paths from browser
    safe_name = secure_filename(file.filename) or "scraped_jobs.txt"
    save_path = os.path.abspath(os.path.join(scanned_dir, safe_name))
    print(f"[DEBUG] Saving txt file to: {save_path}")
    
    with open(save_path, 'w', encoding='utf-8') as f:
        f.write(content)

    return jsonify({
        "success": True,
        "filename": safe_name,
        "char_count": len(content),
        "preview": content[:1000]
    })


# ── Set sample email template ──
@app.route('/api/sample-email', methods=['POST'])
def set_sample_email():
    # No operation: frontend saves it now.
    return jsonify({"success": True})


# ── Analyze jobs with AI (OpenRouter) ──
@app.route('/api/analyze', methods=['POST'])
def analyze_jobs():
    data = request.json or {}
    gemini_api_key = data.get("gemini_api_key", "").strip()

    if not gemini_api_key:
        return jsonify({"success": False, "error": "No Gemini API key provided. Please enter your API key in Settings."}), 400

    txt_content = data.get("txt_content", "")
    sample_email = data.get("sample_email", "Professional email")
    resume_text = data.get("resume_text", "")
    user_name = data.get("user_name", "")

    # Fallback to load from disk if session is empty (e.g. server reset)
    import glob
    if not txt_content:
        scanned_dir = "/tmp/scanned_jobs" if os.environ.get("VERCEL") else "scanned_jobs"
        txt_files = glob.glob(os.path.join(scanned_dir, "*.txt"))
        if txt_files:
            latest = max(txt_files, key=os.path.getmtime)
            with open(latest, 'r', encoding='utf-8', errors='ignore') as f:
                txt_content = f.read()

    if not resume_text:
        uploads_dir = UPLOAD_DIR
        resume_files = glob.glob(os.path.join(uploads_dir, "resume_*.*"))
        if resume_files:
            latest = max(resume_files, key=os.path.getmtime)
            ext = latest.lower().rsplit('.', 1)[-1]
            if ext == 'pdf':
                resume_text = extract_text_from_pdf(latest)
            elif ext == 'txt':
                with open(latest, 'r', encoding='utf-8', errors='ignore') as f:
                    resume_text = f.read()
            elif ext in ['docx', 'doc'] and DOCX_AVAILABLE:
                try:
                    from docx import Document
                    doc = Document(latest)
                    resume_text = '\n'.join([p.text for p in doc.paragraphs])
                except:
                    pass

    if not txt_content:
        return jsonify({"success": False, "error": "No job text to analyze. Upload a .txt file first."}), 400

    user_sample_email_safe = sample_email if sample_email.strip() else "Professional email"
    user_name_display = user_name if user_name else "[Your Name]"

    PROMPT = f"""
You are a highly skilled AI assistant specialized in analyzing job postings and drafting professional job application emails.

Your input text may contain multiple, unstructured job postings scraped from LinkedIn or other sources.

Your tasks:

1. Identify ALL distinct job postings in the input text.
2. FILTER OUT any job that:
   - Does NOT provide a valid apply email
   - Is located OUTSIDE India
3. For each remaining job:
   - Extract ONLY factual information explicitly present in the text
   - Generate a professional, polite, concise email draft by FOLLOWING the STYLE and STRUCTURE of the template below
   - Ensure emails are human-like, coherent, and well-formatted

Strict JSON output schema:

{{
  "jobs": [
    {{
      "job_title": string,
      "company": string,
      "apply_email": string,
      "job_type": "Internship" | "Full-time" | "Contract" | "Part-time" | "Unknown",
      "location": string,
      "skills": string or null,
      "jd_summary": string,         # 1-2 sentences summarizing role, tech stack, and expectations
      "description": string,        # Important job details formatted as: 
                                    # 📌 About Company: (brief company info if available)
                                    # 💼 Key Responsibilities: (main tasks/duties)
                                    # ✅ Requirements: (qualifications, experience needed)
                                    # 💰 Compensation: (salary/stipend if mentioned)
                                    # ⚠️ Important Points: (deadline, work hours, special notes, red flags if any)
      "email_subject": string,      # Clear, professional subject, e.g., "Application for <Job Title> role"
      "email_body_draft": string    # Polished email, max 2 short paragraphs + closing
    }}
  ]
}}

Rules for `email_body_draft`:
- Style reference template (DO NOT COPY TEXT):

\"\"\"
{user_sample_email_safe}
\"\"\"
- Preserve tone and structure
- Lightly customize for each job using job title, skills, and JD summary
- Keep paragraphs short and readable
- Do NOT exaggerate, invent skills, or fabricate experience
- Ensure proper grammar and professional formatting
- IMPORTANT: Use explicit `\\n\\n` characters to separate paragraphs.
- IMPORTANT: End email with "Yours sincerely," followed by the applicant's name: "{user_name_display}"

Additional instructions:
- Output JSON ONLY
- Do NOT include explanations, comments, or non-job content
- Ensure all extracted jobs are unique
- Validate that `apply_email` looks legitimate (contains "@" and domain)
- Include only jobs in India

TEXT TO ANALYZE:
"""

    def call_ai(prompt_text, retries=4):
        """Call Gemini API natively using google.generativeai."""
        import time
        import google.generativeai as genai
        for attempt in range(retries + 1):
            try:
                genai.configure(api_key=gemini_api_key)
                model = genai.GenerativeModel("gemini-2.5-flash")
                
                response = model.generate_content(
                    prompt_text,
                    generation_config=genai.GenerationConfig(temperature=0),
                )
                
                return response.text
            except Exception as e:
                err_str = str(e).lower()
                if "429" in err_str or "quota" in err_str or "rate limit" in err_str:
                    if attempt < retries:
                        wait = 15 * (attempt + 1)
                        print(f"[Wait] Rate limited. Waiting {wait}s before retry {attempt + 1}/{retries}...")
                        time.sleep(wait)
                        continue
                    raise Exception("Google Gemini Free Tier Rate Limit hit. Please wait a minute and try again.")
                
                print(f"[Wait] API error: {e}. Retry {attempt + 1}/{retries}...")
                if attempt < retries:
                    time.sleep(5 * (attempt + 1))
                    continue
                raise e

    def extract_json(text):
        """Robustly extract and repair JSON from AI response text."""
        import re as _re
        text = text.strip()
        # Strip markdown code fences
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:])
            if text.rstrip().endswith("```"):
                text = text.rstrip()[:-3].rstrip()
        # Find the JSON object
        start = text.find("{")
        end = text.rfind("}") + 1
        if start != -1 and end > start:
            text = text[start:end]
        # Fix trailing commas before } or ]
        text = _re.sub(r',\s*([}\]])', r'\1', text)
        # Try parsing directly first
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        # Try fixing truncated JSON by closing unclosed brackets
        open_braces = text.count('{') - text.count('}')
        open_brackets = text.count('[') - text.count(']')
        # Remove any trailing incomplete entry (after last complete object)
        last_complete = text.rfind('},')
        if last_complete == -1:
            last_complete = text.rfind('}')
        if last_complete > 0:
            text = text[:last_complete + 1]
        text += ']' * max(0, open_brackets) + '}' * max(0, open_braces)
        text = _re.sub(r',\s*([}\]])', r'\1', text)
        return json.loads(text)

    try:
        response_text = call_ai(PROMPT + "\n\n" + txt_content)
        if not response_text:
            return jsonify({"success": False, "error": "AI model returned an empty response. It may be overloaded. Please try again."}), 500
        
        response_text = response_text.strip()
        parsed_output = extract_json(response_text)
        _raw_jobs = parsed_output.get("jobs", [])
        
        # Programmatically clean and filter out any job without a valid-looking email
        jobs = []
        for j in _raw_jobs:
            email_raw = str(j.get("apply_email", "")).strip()
            if "@" in email_raw:
                # Basic cleanup: extract just the email part
                for part in email_raw.split():
                    if "@" in part:
                        j["apply_email"] = part.strip("<>().:,").lower()
                        jobs.append(j)
                        break

        # Add job_id and match_score
        for idx, job in enumerate(jobs, start=1):
            job["job_id"] = idx
            job["match_score"] = 0

        # Score jobs against resume if available
        if resume_text and jobs:
            try:
                score_prompt = f"""
Score each job from 0-100 based on how well it matches this resume.

RESUME:
{resume_text[:2500]}

JOBS:
{json.dumps([{"job_id": j["job_id"], "job_title": j["job_title"], "skills": j.get("skills"), "jd_summary": j.get("jd_summary")} for j in jobs], indent=2)[:3000]}

Output JSON only: {{"scores": [{{"job_id": 1, "score": 85}}, ...]}}
Sort by score descending. Score ALL jobs.
"""
                score_text = call_ai(score_prompt)
                score_result = extract_json(score_text.strip() if score_text else "{}")
                scores = {s["job_id"]: s["score"] for s in score_result.get("scores", [])}

                for job in jobs:
                    job["match_score"] = scores.get(job["job_id"], 0)

                jobs.sort(key=lambda j: j["match_score"], reverse=True)

                for idx, job in enumerate(jobs, start=1):
                    job["job_id"] = idx

            except Exception as e:
                print(f"Scoring error: {e}")

                # Removed global session list of jobs

        return jsonify({
            "success": True,
            "job_count": len(jobs),
            "jobs": jobs
        })

    except json.JSONDecodeError as e:
        return jsonify({"success": False, "error": f"Failed to parse AI response as JSON: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"success": False, "error": f"AI analysis failed: {str(e)}"}), 500


    # Rendered entirely by local frontend storage now
    return jsonify({
        "success": True,
        "jobs": request.json.get("jobs", []) if request.is_json else []
    })


# ── Get tracker data ──
@app.route('/api/tracker', methods=['GET'])
def get_tracker():
    return jsonify({
        "success": True,
        "sent_emails": [],
        "tracker": []
    })


# ── Send email via SMTP ──
@app.route('/api/send-email', methods=['POST'])
def send_email():
    data = request.json or {}

    to_email = data.get("to", "")
    subject = data.get("subject", "")
    body = data.get("body", "")
    job_id = data.get("job_id", "")
    job_title = data.get("job_title", "")
    company = data.get("company", "")

    sender_email = data.get("sender_email", "")
    sender_password = data.get("sender_password", "")

    if not sender_email or not sender_password:
        return jsonify({"success": False, "error": "Sender email credentials not configured. Go to Settings."}), 400
    if not to_email:
        return jsonify({"success": False, "error": "Recipient email is required"}), 400
    if not subject or not body:
        return jsonify({"success": False, "error": "Email subject and body are required"}), 400

    # Check if already sent (Handled by Supabase DB via Frontend)

    # Build email message
    msg = EmailMessage()
    msg["From"] = sender_email
    msg["To"] = to_email
    msg["Subject"] = subject
    msg.set_content(body)

    resume_path = data.get("resume_path", "").strip()
    print(f"[DEBUG send-email] resume_path received: {repr(resume_path)}")
    
    # Fallback: if path doesn't exist, try to find the most recent uploaded resume
    if not resume_path or not os.path.exists(resume_path):
        import glob as _glob
        candidates = _glob.glob(os.path.join(UPLOAD_DIR, "resume_*.*"))
        if candidates:
            resume_path = max(candidates, key=os.path.getmtime)
            print(f"[DEBUG send-email] Falling back to latest resume on disk: {resume_path}")
        else:
            print(f"[DEBUG send-email] No resume found in {UPLOAD_DIR}")

    if resume_path and os.path.exists(resume_path):
        print(f"[DEBUG send-email] Attaching resume: {resume_path}")
        with open(resume_path, 'rb') as f:
            resume_data = f.read()
            resume_name = data.get("resume_original_name", "") or os.path.basename(resume_path)
            ext = resume_name.rsplit('.', 1)[-1].lower() if '.' in resume_name else 'pdf'
            mime_map = {
                'pdf': ('application', 'pdf'),
                'docx': ('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document'),
                'doc': ('application', 'msword'),
                'txt': ('text', 'plain'),
                'png': ('image', 'png'),
                'jpg': ('image', 'jpeg'),
                'jpeg': ('image', 'jpeg'),
            }
            maintype, subtype = mime_map.get(ext, ('application', 'octet-stream'))
            msg.add_attachment(
                resume_data,
                maintype=maintype,
                subtype=subtype,
                filename=resume_name
            )
    else:
        print(f"[DEBUG send-email] Skipping attachment — file not found at: {repr(resume_path)}")


    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)

        return jsonify({"success": True, "message": f"Email sent to {to_email}", "sheet_url": None})

    except smtplib.SMTPAuthenticationError:
        return jsonify({"success": False, "error": "Gmail authentication failed. Check your email and app password."}), 401
    except Exception as e:
        return jsonify({"success": False, "error": f"Failed to send: {str(e)}"}), 500


# ── Get / init Google Sheet for user ──
@app.route('/api/sheet-url', methods=['GET'])
def sheet_url_endpoint():
    """Return the Google Sheet URL for the current user."""
    user_email = session.get("sender_email", "")
    if not user_email:
        return jsonify({"success": False, "error": "No user logged in"}), 400

    # Cache removed from backend, managed by frontend args if needed
    if not SHEETS_AVAILABLE:
        return jsonify({"success": False, "error": "Google Sheets not configured", "sheets_available": False})

    try:
        url = get_sheet_url(user_email)
        if url:
            return jsonify({"success": True, "sheet_url": url, "sheets_available": True})
        else:
            return jsonify({"success": False, "error": "No sheet found. Send an email to create one.", "sheets_available": True})
    except Exception as e:
        return jsonify({"success": False, "error": str(e), "sheets_available": True}), 500


@app.route('/api/init-sheet', methods=['POST'])
def init_sheet():
    """Manually create/init the Google Sheet for the current user."""
    user_email = session.get("sender_email", "")
    if not user_email:
        return jsonify({"success": False, "error": "No user logged in"}), 400
    if not SHEETS_AVAILABLE:
        return jsonify({"success": False, "error": "Google Sheets not configured. Add service_account.json"}), 400

    try:
        _, url, is_new = get_or_create_sheet(user_email)
        return jsonify({"success": True, "sheet_url": url, "is_new": is_new})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/api/sheet-data', methods=['GET'])
def sheet_data_endpoint():
    """Get tracker data from the user's Google Sheet."""
    user_email = session.get("sender_email", "")
    if not user_email:
        return jsonify({"success": False, "error": "No user logged in"}), 400
    if not SHEETS_AVAILABLE:
        return jsonify({"success": False, "error": "Google Sheets not configured"}), 400

    try:
        data = get_sheet_data(user_email)
        return jsonify({"success": True, "records": data, "count": len(data)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# ── Session info ──
@app.route('/api/session', methods=['GET'])
def get_session():
    # Now totally handled by localStorage in the UI directly, so we just return config.
    return jsonify({
        "sheets_available": SHEETS_AVAILABLE,
    })


# =========================
# ADMIN – All user sheets
# =========================

@app.route('/api/admin/sheets', methods=['GET'])
def admin_list_sheets():
    """List ALL user Google Sheets (admin view)."""
    if not SHEETS_AVAILABLE:
        return jsonify({"success": False, "error": "Google Sheets not configured"}), 400
    try:
        sheets = list_all_user_sheets()
        return jsonify({"success": True, "sheets": sheets, "total_users": len(sheets)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route('/admin')
def admin_page():
    return send_from_directory('ui', 'admin.html')


# =========================
# MAIN
# =========================
if __name__ == '__main__':
    print("\n" + "="*50)
    print("  [x] RoleMatch AI Server")
    print("  --> http://localhost:5000")
    print("="*50 + "\n")
    app.run(debug=True, use_reloader=False, port=5000)
